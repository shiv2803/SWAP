import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Dictionary of all extracted source file names and their exact code contents
ALL_CODES = {
    "safety_rules.py": """import time

class SafetyController:
    def __init__(self):
        self.MIN_DWELL_TIME = 3.0
        self.CONFIDENCE_THRESHOLD = 0.75
        self.MAX_LORA_DUTY_CYCLE = 0.01
        self.TELEMETRY_STALE_AGE = 2.5
        
        self.current_protocol = 0
        self.last_switch_time = time.time()
        self.last_telemetry_time = time.time()
        
        self.lora_tx_time_total = 0.0
        self.lora_window_start = time.time()

    def update_telemetry_timestamp(self):
        self.last_telemetry_time = time.time()

    def log_lora_tx(self, tx_duration_seconds):
        self.lora_tx_time_total += tx_duration_seconds

    def _is_telemetry_stale(self, now):
        return (now - self.last_telemetry_time) > self.TELEMETRY_STALE_AGE

    def _check_lora_duty_cycle_available(self, now):
        window_elapsed = now - self.lora_window_start
        if window_elapsed >= 60.0:
            self.lora_tx_time_total = 0.0
            self.lora_window_start = now
            window_elapsed = 0.1
            
        current_duty_cycle = self.lora_tx_time_total / window_elapsed
        return current_duty_cycle < self.MAX_LORA_DUTY_CYCLE

    def evaluate(self, ml_target, ml_confidence):
        now = time.time()
        if self._is_telemetry_stale(now):
            print("[SAFETY] Stale telemetry. Hard fallback to LoRa (2).")
            return self._execute_switch(2, now)

        if (now - self.last_switch_time) < self.MIN_DWELL_TIME:
            return self.current_protocol

        if ml_confidence < self.CONFIDENCE_THRESHOLD:
            print(f"[SAFETY] Low confidence ({ml_confidence:.2f}). Maintaining {self.current_protocol}.")
            return self.current_protocol

        if ml_target == 2 and not self._check_lora_duty_cycle_available(now):
            print("[SAFETY] LoRa duty cycle exceeded. Hard fallback to Wi-Fi (0).")
            return self._execute_switch(0, now)

        if ml_target != self.current_protocol:
            print(f"[SAFETY] ML switch approved: {self.current_protocol} -> {ml_target}")
            return self._execute_switch(ml_target, now)

        return self.current_protocol

    _execute_switch = lambda self, p, t: setattr(self, 'current_protocol', p) or setattr(self, 'last_switch_time', t) or self.current_protocol

if __name__ == "__main__":
    safety = SafetyController()
    print("Safety module loaded successfully.")""",

    "laptop_training/train_random_forest.py": """import os
import pandas as pd
import numpy as np
from sklearn.ensemble import RandomForestClassifier
from sklearn.model_selection import train_test_split
from sklearn.metrics import classification_report, confusion_matrix
import joblib

CSV_FILENAME = "swap_rf_data.csv"
MODEL_FILENAME = "swap_random_forest.joblib"

FEATURE_COLS = ["wifi_rssi", "ble_rssi", "lora_rssi", "lora_snr", "pdr", "rtt", "retries"]
TARGET_COL = "best_protocol"

def generate_dummy_data_if_missing():
    if os.path.exists(CSV_FILENAME):
        return
    np.random.seed(42)
    data = {
        "wifi_rssi": np.random.randint(-90, -40, 500),
        "ble_rssi": np.random.randint(-100, -50, 500),
        "lora_rssi": np.random.randint(-120, -80, 500),
        "lora_snr": np.random.uniform(-10.0, 10.0, 500),
        "pdr": np.random.uniform(0.5, 1.0, 500),
        "rtt": np.random.uniform(10.0, 500.0, 500),
        "retries": np.random.randint(0, 5, 500),
        "best_protocol": np.random.choice([0, 1, 2], 500)
    }
    pd.DataFrame(data).to_csv(CSV_FILENAME, index=False)

def main():
    generate_dummy_data_if_missing()
    df = pd.read_csv(CSV_FILENAME)
    X, y = df[FEATURE_COLS], df[TARGET_COL]
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.25, random_state=42, stratify=y)
    
    rf = RandomForestClassifier(n_estimators=100, max_depth=8, min_samples_leaf=5, class_weight="balanced", random_state=42, n_jobs=-1)
    rf.fit(X_train, y_train)
    joblib.dump(rf, MODEL_FILENAME)
    with open("feature_sequence.txt", "w") as f:
        f.write(",".join(FEATURE_COLS))

if __name__ == "__main__":
    main()""",

    "laptop_training/generate_dummy_data.py": """import os
import pandas as pd
import numpy as np

CSV_FILENAME = "swap_rf_data.csv"

def generate_dummy_data():
    np.random.seed(42)
    wifi_data = {"wifi_rssi": np.random.randint(-70, -40, 200), "ble_rssi": np.random.randint(-100, -60, 200), "lora_rssi": np.random.randint(-120, -90, 200), "lora_snr": np.random.uniform(-5.0, 5.0, 200), "pdr": np.random.uniform(0.9, 1.0, 200), "rtt": np.random.uniform(10.0, 50.0, 200), "retries": np.random.randint(0, 2, 200), "best_protocol": [0] * 200}
    ble_data = {"wifi_rssi": np.random.randint(-95, -75, 200), "ble_rssi": np.random.randint(-75, -50, 200), "lora_rssi": np.random.randint(-120, -90, 200), "lora_snr": np.random.uniform(-5.0, 5.0, 200), "pdr": np.random.uniform(0.8, 1.0, 200), "rtt": np.random.uniform(20.0, 100.0, 200), "retries": np.random.randint(0, 3, 200), "best_protocol": [1] * 200}
    lora_data = {"wifi_rssi": np.random.randint(-100, -85, 200), "ble_rssi": np.random.randint(-100, -85, 200), "lora_rssi": np.random.randint(-110, -70, 200), "lora_snr": np.random.uniform(0.0, 12.0, 200), "pdr": np.random.uniform(0.4, 0.8, 200), "rtt": np.random.uniform(100.0, 1000.0, 200), "retries": np.random.randint(2, 6, 200), "best_protocol": [2] * 200}
    
    df = pd.concat([pd.DataFrame(wifi_data), pd.DataFrame(ble_data), pd.DataFrame(lora_data)]).sample(frac=1).reset_index(drop=True)
    df.to_csv(CSV_FILENAME, index=False)

if __name__ == "__main__":
    generate_dummy_data()""",

    "swap_backend/__init__.py": '\"\"\"SWAP prototype backend package.\"\"\"',

    "swap_backend/common.py": """import csv
import logging
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any, Dict, Optional

logger = logging.getLogger("swap.common")

PROTOCOL_WIFI = 0
PROTOCOL_BLE = 1
PROTOCOL_LORA = 2
PROTOCOL_NAMES = {
    PROTOCOL_WIFI: "WIFI",
    PROTOCOL_BLE: "BLE",
    PROTOCOL_LORA: "LORA",
}

CSV_FIELDS = [
    "recv_ts", "node", "ts_ms", "active_protocol",
    "wifi_rssi", "wifi_loss", "ble_rssi",
    "lora_rssi", "lora_snr", "lora_loss", "rtt_ms"
]

@dataclass
class TelemetryRecord:
    recv_ts: float
    node: str
    ts_ms: int
    active_protocol: int
    wifi_rssi: float
    wifi_loss: float
    ble_rssi: float
    lora_rssi: float
    lora_snr: float
    lora_loss: float
    rtt_ms: int

    @staticmethod
    def from_json(raw: Dict[str, Any], recv_ts: float) -> Optional["TelemetryRecord"]:
        try:
            return TelemetryRecord(
                recv_ts=recv_ts,
                node=str(raw["node"]).lower(),
                ts_ms=int(raw["ts_ms"]),
                active_protocol=int(raw["active_protocol"]),
                wifi_rssi=float(raw.get("wifi_rssi", -999.0)),
                wifi_loss=float(raw.get("wifi_loss", 1.0)),
                ble_rssi=float(raw.get("ble_rssi", -999.0)),
                lora_rssi=float(raw.get("lora_rssi", -999.0)),
                lora_snr=float(raw.get("lora_snr", -999.0)),
                lora_loss=float(raw.get("lora_loss", 1.0)),
                rtt_ms=int(raw.get("rtt_ms", 0)),
            )
        except (KeyError, TypeError, ValueError):
            return None

class CsvLogger:
    def __init__(self, path: str):
        self._path = Path(path)
        self._path.parent.mkdir(parents=True, exist_ok=True)
        if not self._path.exists() or self._path.stat().st_size == 0:
            with self._path.open("w", newline="", encoding="utf-8") as f:
                csv.DictWriter(f, fieldnames=CSV_FIELDS).writeheader()

    def append(self, record: TelemetryRecord) -> None:
        with self._path.open("a", newline="", encoding="utf-8") as f:
            csv.DictWriter(f, fieldnames=CSV_FIELDS).writerow(asdict(record))""",

    "swap_backend/simulator.py": """import asyncio
import math
import random
import time
from dataclasses import dataclass
from typing import Dict, Optional
from .common import PROTOCOL_BLE, PROTOCOL_LORA, PROTOCOL_WIFI, CsvLogger, TelemetryRecord

@dataclass
class _NodePhase:
    wifi_rssi: float
    wifi_loss: float
    ble_rssi: float
    lora_rssi: float
    lora_snr: float
    lora_loss: float

class SimulatorCommandSender:
    def __init__(self, source):
        self._source = source
    async def force_protocol(self, node: str, protocol: int) -> None:
        self._source.force_protocol(node=node, protocol=protocol)

class SimulatorTelemetrySource:
    def __init__(self, csv_logger: CsvLogger, interval_s: float = 0.25):
        self._csv_logger = csv_logger
        self._interval_s = interval_s
        self._forced_protocol: Dict[str, int] = {}
        self._sender = None

    def attach_sender(self, sender) -> None:
        self._sender = sender

    def command_sender(self):
        return self._sender

    def force_protocol(self, node: str, protocol: int) -> None:
        self._forced_protocol[node] = protocol

    async def run(self, out_queue) -> None:
        await asyncio.gather(
            self._node_loop("a", 0.0, out_queue),
            self._node_loop("b", 3.2, out_queue),
        )

    async def _node_loop(self, node: str, phase_offset: float, out_queue) -> None:
        while True:
            now = time.time()
            cycle = (math.sin((now + phase_offset) / 18.0) + 1.0) / 2.0
            phase = _NodePhase(-60.0 - 25.0 * cycle, 0.02 + 0.38 * cycle, -70.0 - 18.0 * cycle, -90.0 - 7.0 * cycle, 7.0 - 4.0 * cycle, 0.01 + 0.07 * cycle)
            active_protocol = self._forced_protocol.get(node, (0 if phase.wifi_rssi > -75.0 and phase.wifi_loss < 0.10 else (1 if phase.ble_rssi > -85.0 else 2)))
            rtt_ms = int(max(1.0, (7.0 if active_protocol == 0 else (25.0 if active_protocol == 1 else 220.0)) + random.uniform(-3.0, 3.0)))
            
            record = TelemetryRecord(now, node, int(now * 1000), active_protocol, phase.wifi_rssi, phase.wifi_loss, phase.ble_rssi, phase.lora_rssi, phase.lora_snr, phase.lora_loss, rtt_ms)
            self._csv_logger.append(record)
            await out_queue.put(record)
            await asyncio.sleep(self._interval_s)""",

    "swap_backend/telemetry_link.py": """import asyncio
import json
import logging
import os
import time
from .common import CsvLogger, TelemetryRecord
from .simulator import SimulatorCommandSender, SimulatorTelemetrySource

logger = logging.getLogger("swap.telemetry")

def create_source_from_env():
    csv_logger = CsvLogger(os.environ.get("SWAP_CSV_LOG", "telemetry_log.csv"))
    sim_source = SimulatorTelemetrySource(csv_logger=csv_logger)
    sim_source.attach_sender(SimulatorCommandSender(sim_source))
    return sim_source""",

    "swap_backend/link_quality_model.py": """import logging
from collections import deque
from dataclasses import dataclass, field
from pathlib import Path
from typing import Deque, Dict, List, Optional
import joblib
import numpy as np
from .common import PROTOCOL_BLE, PROTOCOL_LORA, PROTOCOL_WIFI, TelemetryRecord

logger = logging.getLogger("swap.model")
MODEL_PATH = Path("models") / "link_quality_model.joblib"
HYSTERESIS_MARGIN = 3.5
SOFTMAX_TEMPERATURE = 10.0
LQI_SCORE_MIN, LQI_SCORE_MAX = -150.0, -30.0

@dataclass
class Decision:
    protocol: int
    confidence: float
    source: str
    raw_scores: Dict[str, float] = field(default_factory=dict)
    probabilities: Dict[int, float] = field(default_factory=dict)
    lqi: int = 0
    hysteresis_guarded: bool = False
    best_candidate: int = 0

class FeatureWindow:
    def __init__(self, window_size: int):
        self._window: Deque[TelemetryRecord] = deque(maxlen=window_size)
    def push(self, record: TelemetryRecord) -> None:
        self._window.append(record)
    def is_ready(self, min_samples: int) -> bool:
        return len(self._window) >= min_samples
    def to_feature_vector(self) -> np.ndarray:
        return np.array([[r.wifi_rssi, r.wifi_loss, r.ble_rssi, r.lora_rssi, r.lora_snr, r.lora_loss, r.rtt_ms] for r in self._window]).mean(axis=0).reshape(1, -1)
    def means_dict(self) -> Dict[str, float]:
        m = np.array([[r.wifi_rssi, r.wifi_loss, r.ble_rssi, r.lora_rssi, r.lora_snr, r.lora_loss, r.rtt_ms] for r in self._window]).mean(axis=0)
        return {"wifi_rssi": float(m[0]), "wifi_loss": float(m[1]), "ble_rssi": float(m[2]), "lora_rssi": float(m[3]), "lora_snr": float(m[4]), "lora_loss": float(m[5]), "rtt_ms": float(m[6])}
    def latest_active_protocol(self) -> int:
        return self._window[-1].active_protocol if self._window else PROTOCOL_WIFI

class LinkQualityModel:
    def __init__(self, model_path: Path = MODEL_PATH, window_size: int = 10, min_samples: int = 3):
        self._model = joblib.load(model_path) if model_path.exists() else None
        self._window_size, self._min_samples, self._windows = window_size, min_samples, {}
    def _window_for(self, node_id: str) -> FeatureWindow:
        if node_id not in self._windows:
            self._windows[node_id] = FeatureWindow(self._window_size)
        return self._windows[node_id]
    def observe(self, record: TelemetryRecord) -> Optional[Decision]:
        w = self._window_for(record.node)
        w.push(record)
        return self._predict_from_window(w) if w.is_ready(self._min_samples) else None
    def recommend_for_node(self, node_id: str) -> Optional[Decision]:
        w = self._windows.get(node_id)
        return self._predict_from_window(w) if w and w.is_ready(self._min_samples) else None
    def _predict_from_window(self, window: FeatureWindow) -> Decision:
        m = window.means_dict()
        scores = {PROTOCOL_WIFI: m["wifi_rssi"] - (m["wifi_loss"] * 120.0) - (m["rtt_ms"] * 0.1), PROTOCOL_BLE: m["ble_rssi"] - (m["wifi_loss"] * 90.0) - (m["rtt_ms"] * 0.15), PROTOCOL_LORA: m["lora_rssi"] + (m["lora_snr"] * 2.0) - (m["lora_loss"] * 160.0) - (m["rtt_ms"] * 0.02)}
        best_candidate = max(scores, key=scores.get)
        if self._model:
            fv = window.to_feature_vector()
            candidate, confidence, source = int(self._model.predict(fv)[0]), float(np.max(self._model.predict_proba(fv)[0])), "model"
        else:
            candidate, confidence, source = (PROTOCOL_WIFI if m["wifi_rssi"] > -75.0 and m["wifi_loss"] < 0.10 else (PROTOCOL_BLE if m["ble_rssi"] > -85.0 else PROTOCOL_LORA)), 1.0, "rule_based"
        return Decision(candidate, confidence, source, {"wifi": round(scores[PROTOCOL_WIFI], 2), "ble": round(scores[PROTOCOL_BLE], 2), "lora": round(scores[PROTOCOL_LORA], 2)}, {k: float(v) for k, v in scores.items()}, int(max(0, min(100, round(((scores[candidate] - LQI_SCORE_MIN) / (LQI_SCORE_MAX - LQI_SCORE_MIN)) * 100.0)))), False, best_candidate)""",

    "swap_backend/app.py": """import asyncio
import contextlib
import logging
from collections import deque
from dataclasses import asdict
from typing import Deque, Dict, List, Optional
from fastapi import FastAPI, HTTPException, Query, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, Field
from .common import PROTOCOL_NAMES, TelemetryRecord
from .link_quality_model import Decision, LinkQualityModel
from .telemetry_link import create_source_from_env

logging.basicConfig(level=logging.INFO)
app = FastAPI(title="SWAP Prototype API", version="1.0.0")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])

telemetry_queue = asyncio.Queue()
telemetry_source = create_source_from_env()
command_sender = telemetry_source.command_sender()
model = LinkQualityModel()
recent_records = deque(maxlen=500)
latest_record_by_node, latest_decision_by_node, ws_clients = {}, {}, []
ingest_task = None

@app.on_event("startup")
async def on_startup():
    global ingest_task
    ingest_task = asyncio.create_task(_ingest_loop())

@app.on_event("shutdown")
async def on_shutdown():
    if ingest_task:
        ingest_task.cancel()
        with contextlib.suppress(asyncio.CancelledError):
            await ingest_task

@app.get("/health")
def health(): return {"status": "ok"}

@app.get("/status")
def get_status():
    return {"nodes": {node: {"active_protocol_reported": r.active_protocol, "active_protocol_name": PROTOCOL_NAMES.get(r.active_protocol), "recommended_protocol": d.protocol if d else None} for node in ("a", "b") if (r := latest_record_by_node.get(node))}}

@app.websocket("/ws/live")
async def live_stream(websocket: WebSocket):
    await websocket.accept()
    ws_clients.append(websocket)
    try:
        while True: await websocket.receive_text()
    except WebSocketDisconnect: pass
    finally:
        if websocket in ws_clients: ws_clients.remove(websocket)

async def _ingest_loop():
    source_task = asyncio.create_task(telemetry_source.run(telemetry_queue))
    try:
        while True:
            record = await telemetry_queue.get()
            recent_records.append(record)
            latest_record_by_node[record.node] = record
            decision = model.observe(record)
            if decision: latest_decision_by_node[record.node] = decision
    finally:
        source_task.cancel()
        with contextlib.suppress(asyncio.CancelledError): await source_task"""
}

def create_code_document():
    doc = Document()
    
    # Title
    title = doc.add_heading("SWAP Project — Full Source Code Documentation", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("This document contains all backend modules, offline model training scripts, and safety controller implementations compiled for the SWAP project.")
    
    for filepath, code_content in ALL_CODES.items():
        doc.add_heading(filepath, level=1)
        p = doc.add_paragraph()
        run = p.add_run(code_content)
        run.font.name = "Consolas"
        run.font.size = Pt(9.5)
        doc.add_page_break()
        
    output_filename = "SWAP_All_Project_Codes.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}")

if __name__ == "__main__":
    create_code_document()